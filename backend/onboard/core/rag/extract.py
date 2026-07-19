from dataclasses import dataclass, field

from docx import Document as DocxDocument
from pypdf import PdfReader

from onboard.core.common.exceptions import ValidationError


@dataclass
class ExtractedPage:
    page_number: int
    text: str
    section_title: str | None = None


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage] = field(default_factory=list)
    page_count: int = 0

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


def extract_document(content: bytes, file_type: str) -> ExtractedDocument:
    """Extract text from an uploaded Doc Pack file (PDF / DOCX / TXT / MD)."""
    normalized = file_type.lower().lstrip(".")
    if normalized == "markdown":
        normalized = "md"

    if normalized == "pdf":
        return _extract_pdf(content)
    if normalized == "docx":
        return _extract_docx(content)
    if normalized in {"txt", "md"}:
        return _extract_plain(content)
    raise ValidationError(f"Unsupported file type: {file_type}")


def _extract_pdf(content: bytes) -> ExtractedDocument:
    from io import BytesIO

    page_count = 0
    pages: list[ExtractedPage] = []
    try:
        reader = PdfReader(BytesIO(content))
        page_count = len(reader.pages)
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(ExtractedPage(page_number=index, text=text))
    except ValidationError:
        raise
    except Exception:
        pages = []

    # pypdf returns empty text for some perfectly valid PDFs (unusual encodings, exotic
    # generators). pdfminer parses the content streams itself, so fall back to it before
    # declaring the document unreadable.
    if not pages:
        pages, fallback_page_count = _extract_pdf_pdfminer(content)
        page_count = page_count or fallback_page_count

    if not pages:
        raise ValidationError(
            "PDF has no selectable text — it is likely a scanned or image-only PDF. "
            "Re-export it as a text-based PDF and retry."
        )
    return ExtractedDocument(pages=pages, page_count=page_count or len(pages))


def _extract_pdf_pdfminer(content: bytes) -> tuple[list[ExtractedPage], int]:
    from io import BytesIO

    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer

    pages: list[ExtractedPage] = []
    page_count = 0
    try:
        for index, layout in enumerate(extract_pages(BytesIO(content)), start=1):
            page_count = index
            text = "".join(el.get_text() for el in layout if isinstance(el, LTTextContainer)).strip()
            if text:
                pages.append(ExtractedPage(page_number=index, text=text))
    except Exception:
        return [], page_count
    return pages, page_count


def _extract_docx(content: bytes) -> ExtractedDocument:
    from io import BytesIO

    document = DocxDocument(BytesIO(content))
    pages: list[ExtractedPage] = []
    current_section: str | None = None
    buffer: list[str] = []

    def flush(page_number: int) -> None:
        nonlocal buffer, current_section
        text = "\n".join(buffer).strip()
        if text:
            pages.append(ExtractedPage(page_number=page_number, text=text, section_title=current_section))
        buffer = []

    page_number = 1
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        text = paragraph.text.strip()
        if not text:
            continue
        if "heading" in style_name:
            flush(page_number)
            current_section = text
            page_number += 1
            buffer = [text]
        else:
            buffer.append(text)

    flush(page_number)
    if not pages:
        raise ValidationError("DOCX contained no extractable text")
    return ExtractedDocument(pages=pages, page_count=len(pages))


def _extract_plain(content: bytes) -> ExtractedDocument:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
    text = text.strip()
    if not text:
        raise ValidationError("File contained no extractable text")
    return ExtractedDocument(pages=[ExtractedPage(page_number=1, text=text)], page_count=1)
